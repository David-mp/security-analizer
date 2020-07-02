#Requirements:
#1-Java 11
#2-Having Sonarqube installed (https://www.sonarqube.org/downloads)
#3-Having SonarScanner installed (https://docs.sonarqube.org/latest/analysis/scan/sonarscanner)
#4-Python3

#----------------------------------------
#----------------------------------------

#Setting up the environment

import os

os.system("pip install pylint")
os.system("pip install requests")

#----------------------------------------

#######################################################################
########################    INPUTS     ################################
#######################################################################

#!-!-!-!-! Path to SonarStart.bat
sonar_dir = "C:/Users/ManagingandInnovatio/Desktop/Sonarqube/sonarqube-8.2.0.32929/bin/windows-x86-64"
#!-!-!-!-! Path to sonar-scanner.bat
scanner_dir = "C:/Users/ManagingandInnovatio/Desktop/Sonarqube/sonar-scanner-4.2.0.1873-windows/bin"
#!-!-!-!-! Path to project root directory
project_dir = "C:/Users/ManagingandInnovatio/Desktop/security_analyzer/example_project"

project_name = "example2"
project_key = "example2"
projectVersion = "1.0"


#######################################################################
###################    Starts Sonarqube      ##########################
#######################################################################

import requests
import time

#Starts Sonarqube in a new command window
os.system("start cmd /k " + sonar_dir + "/StartSonar.bat")

#Checks if Sonarqube is running
try:
    r = requests.get(url = "http://localhost:9000")
except Exception as e:
    #Wait some time for Sonarqube to start and try again
    print("Waiting for Sonarqube to start...")
    time.sleep(60)

#######################################################################
###################    Analysis setup      ############################
#######################################################################

#---------Login and authentication check----------------
url = 'http://localhost:9000/api/authentication/login'
login = [("login","admin"), ("password","admin")]
x = requests.post(url, data = login)

url = 'http://localhost:9000/api/authentication/validate'
r = requests.get(url = url)
data = r.json()
print(data)

#---------Creates new project----------------------------
URL = 'http://localhost:9000/api/components/show'
PARAMS = [("component",project_name)]
r = requests.get(url = URL, params = PARAMS)
data = r.json()

#Create the project if it does not exist
if "errors" in data:
    url = 'http://localhost:9000/api/projects/create'
    PARAMS = [("name",project_name), ("project",project_key)]
    x = requests.post(url, data = PARAMS)
    print("-------------------------------")
    print("Creating Sonarqube project...")
    print("-------------------------------")

#----------Launches analysis-----------------------------
analysis_command = scanner_dir + "/sonar-scanner.bat" + " -Dsonar.projectName=" + project_name + " -Dsonar.projectKey=" + project_key + " -Dsonar.login=admin -Dsonar.password=admin " + " -Dsonar.sources=." + " -Dsonar.host.url=http://localhost:9000" + " -Dsonar.projectVersion=" + projectVersion +  " -Dsonar.projectBaseDir=" + project_dir

print("-------------------------------")
print("Analyzing the project...")
print("-------------------------------")

os.system(analysis_command)

#######################################################################
##########    Searchs for Owasp vulnerabilities      ##################
#######################################################################


#!-!-!-!-! 
#project_name = "Python-Proyect-recursive_IDS"
URL = "http://localhost:9000/api/issues/search"
owasp = "owaspTop10"
PARAMS = {'componentKeys':project_name, 'facets':owasp}

try:
    r = requests.get(url = URL, params = PARAMS)
except Exception as e:
    #Wait some time for Sonarqube to start and try again
    time.sleep(30)
    r = requests.get(url = URL, params = PARAMS)

data = r.json()
total = data['total']
issues = data['issues']
facets = data['facets']
owasp_Tags = ["owasp-a1", "owasp-a2", "owasp-a3", "owasp-a4", "owasp-a5", "owasp-a6", "owasp-a7", "owasp-a8", "owasp-a9", "owasp-a10"]

#####################################################
############ Filters issues by Owasp tags ###########
#####################################################

#List of issues encountered of each owasp vulnerability
#"a1": [{issue1}, {issue2}, {issue3}, ...]
filtered_issues = {"a1": [],
          "a2": [],
          "a3": [],
          "a4": [],
          "a5": [],
          "a6": [],
          "a7": [],
          "a8": [],
          "a9": [],
          "a10": []
          }
    
for i in issues:
    tags = i['tags']
    for tag in tags:
        if tag in owasp_Tags:
            #From "owasp-aX" gets "aX"
            key = tag[6:]
            filtered_issues[key].append(i)
            print(tag)

print("Project name: ",project_name)
print("Total issues: ",total)
#print(sonar.sources)
print()
print(facets)

#####################################################
############ Vulnerabilities Description ############
#####################################################

vulnerabilities = {
  "a1": {
    'Name': 'Code injection', 
    'Description': 'The software constructs all or part of a code segment using externally-influenced input from an upstream component, but it does not neutralize or incorrectly neutralizes special elements that could modify the syntax or behavior of the intended code segment.', 
    'Image': 'a1.png', 
    'Possible mitigations': {
      "Phase: Arquitecture and design" : [
        "•  If at all possible, use library calls rather than external processes to recreate the desired functionality.",
        "•  Run your code using the lowest privileges that are required to accomplish the necessary tasks. If possible, create isolated accounts with limited privileges that are only used for a single task. That way, a successful attack will not immediately give the attacker access to the rest of the software or its environment. For example, database applications rarely need to run as the database administrator, especially in day-to-day operations. Specifically, follow the principle of least privilege when creating user accounts to a SQL database. The database users should only have the minimum privileges necessary to use their account. If the requirements of the system indicate that a user can read and modify their own data, then limit their privileges so they cannot read/write others' data. Use the strictest permissions possible on all database objects, such as execute-only for stored procedures.",
        "•  Parameterization: If available, use structured mechanisms that automatically enforce the separation between data and code. These mechanisms may be able to provide the relevant quoting, encoding, and validation automatically, instead of relying on the developer to provide this capability at every point where output is generated. Some languages offer multiple functions that can be used to invoke commands. Where possible, identify any function that invokes a command shell using a single string, and replace it with a function that requires individual arguments. These functions typically perform appropriate quoting and filtering of arguments. For example, in C, the system() function accepts a string that contains the entire command to be executed, whereas execl(), execve(), and others require an array of strings, one for each argument. In Windows, CreateProcess() only accepts one command at a time. In Perl, if system() is provided with an array of arguments, then it will quote each of the arguments."
          ],
      "Phase: Implementation" : [
        '•  Input Validation: Assume all input is malicious. Use a list of acceptable inputs that strictly conform to specifications. Reject any input that does not strictly conform to specifications, or transform it into something that does. When performing input validation, consider all potentially relevant properties, including length, type of input, the full range of acceptable values, missing or extra inputs, syntax, consistency across related fields, and conformance to business rules. As an example of business rule logic, "boat" may be syntactically valid because it only contains alphanumeric characters, but it is not valid if the input is only expected to contain colors such as "red" or "blue."',
        "•  Do not rely exclusively on looking for malicious or malformed inputs. This is likely to miss at least one undesirable input, especially if the code's environment changes. This can give attackers enough room to bypass the intended validation. However, blacklists can be useful for detecting potential attacks or determining which inputs are so malformed that they should be rejected outright.",
        "•  If the program to be executed allows arguments to be specified within an input file or from standard input, then consider using that mode to pass arguments instead of the command line.",
        "•  Output Encoding: While it is risky to use dynamically-generated query strings, code, or commands that mix control and data together, sometimes it may be unavoidable. Properly quote arguments and escape any special characters within those arguments. The most conservative approach is to escape or filter all characters that do not pass an extremely strict whitelist (such as everything that is not alphanumeric or white space). If some special characters are still needed, such as white space, wrap each argument in quotes after the escaping/filtering step. Be careful of argument injection."
      ],
      "Phase: System Configuration" : [
        "•  Assign permissions to the software system that prevents the user from accessing/opening privileged files."]
    },
    'Reference links': ["https://cwe.mitre.org/data/definitions/77.html", "https://cwe.mitre.org/data/definitions/78.html", "https://cwe.mitre.org/data/definitions/88.html", "https://cwe.mitre.org/data/definitions/89.html", "https://cwe.mitre.org/data/definitions/90.html", "https://cwe.mitre.org/data/definitions/91.html", "https://cwe.mitre.org/data/definitions/564.html", "https://cwe.mitre.org/data/definitions/917.html", "https://cwe.mitre.org/data/definitions/943.html"]
  },
  "a2": {
    'Name': 'Broken Authentication',
    'Description': 'When an actor claims to have a given identity, the software does not prove or insufficiently proves that the claim is correct.',
    'Image': 'a2.png',
    'Possible mitigations': {
      "Phase: Arquitecture and design" : [
        "•  Libraries or Frameworks: Use an authentication framework or library such as the OWASP ESAPI Authentication feature.",
        "•  Avoid storing passwords in easily accessible locations.",
        "•  Consider storing cryptographic hashes of passwords as an alternative to storing in plaintext.",
        "•  Use multiple independent authentication schemes, which ensures that -- if one of the methods is compromised -- the system itself is still likely safe from compromise.",
        "•  Invalidate any existing session identifiers prior to authorizing a new user session.",
        '•  When prompting for a password change, force the user to provide the original password in addition to the new password. Do not use "forgotten password" functionality. But if you must, ensure that you are only providing information to the actual user, e.g. by using an email address or challenge question that the legitimate user already provided in the past; do not allow the current user to change this identity information until the correct password has been provided.',
        "•  Make sure that there is throttling on the number of incorrect answers to a security question. Disable the password recovery functionality after a certain (small) number of incorrect guesses. Require that the user properly answers the security question prior to resetting their password and sending the new password to the e-mail address of record."
          ],
      "Phase: Implementation" : [
        "•  Set sessions/credentials expiration date."
      ],
      "Phase: System Configuration" : [
        "•  Enforce SSL use for the login page or any page used to transmit user credentials or other sensitive information. Even if the entire site does not use SSL, it MUST use SSL for login. Additionally, to help prevent phishing attacks, make sure that SSL serves the login page. SSL allows the user to verify the identity of the server to which they are connecting. If the SSL serves login page, the user can be certain they are talking to the proper end system. A phishing attack would typically redirect a user to a site that does not have a valid trusted server certificate issued from an authorized supplier."]
    },
    'Reference links': ["https://cwe.mitre.org/data/definitions/287.html", "https://cwe.mitre.org/data/definitions/256.html", "https://cwe.mitre.org/data/definitions/308.html", "https://cwe.mitre.org/data/definitions/384.html", "https://cwe.mitre.org/data/definitions/522.html", "https://cwe.mitre.org/data/definitions/523.html", "https://cwe.mitre.org/data/definitions/613.html", "https://cwe.mitre.org/data/definitions/620.html", "https://cwe.mitre.org/data/definitions/640.html"]
  },
  "a3": {
    'Name': 'Sensitive Data Exposure',
    'Description': 'The software does not encrypt sensitive or critical information before storage or transmission.', 
    'Image': 'a3.png',
    'Possible mitigations': {
      "Phase: Requirements" : [
        "•  Clearly specify which data or resources are valuable enough that they should be protected by encryption. Require that any transmission or storage of this data/resource should use well-vetted encryption algorithms."
        ],
      "Phase: Arquitecture and design" : [
        "•  Carefully manage and protect cryptographic keys (see CWE-320). If the keys can be guessed or stolen, then the strength of the cryptography itself is irrelevant.",
        "•  Libraries or Frameworks: When there is a need to store or transmit sensitive data, use strong, up-to-date cryptographic algorithms to encrypt that data. Select a well-vetted algorithm that is currently considered to be strong by experts in the field, and use well-tested implementations. As with all cryptographic mechanisms, the source code should be available for analysis. /n Do not develop custom or private cryptographic algorithms. They will likely be exposed to attacks that are well-understood by cryptographers. Reverse engineering techniques are mature. If the algorithm can be compromised if attackers find out how it works, then it is especially weak. /nPeriodically ensure that the cryptography has not become obsolete. Some older algorithms, once thought to require a billion years of computing time, can now be broken in days or hours. This includes MD4, MD5, SHA1, DES, and other algorithms that were once regarded as strong."
        ],
      "Phase: Implementation" : [
        "•  Certificates should be carefully managed and checked to assure that data are encrypted with the intended owner's public key. If certificate pinning is being used, ensure that all relevant properties of the certificate are fully validated before the certificate is pinned, including the hostname.",
        "•  When using web applications with SSL, use SSL for the entire session from login to logout, not just for the initial login page."
      ],
      "Phase: System Configuration" : [
        "•  Access control permissions should be set to prevent reading/writing of sensitive files inside/outside of the FTP directory."
        ]
    },
    'Reference links': ["https://cwe.mitre.org/data/definitions/220.html", "https://cwe.mitre.org/data/definitions/295.html", "https://cwe.mitre.org/data/definitions/311.html", "https://cwe.mitre.org/data/definitions/312.html", "https://cwe.mitre.org/data/definitions/319.html", "https://cwe.mitre.org/data/definitions/325.html", "https://cwe.mitre.org/data/definitions/326.html", "https://cwe.mitre.org/data/definitions/327.html", "https://cwe.mitre.org/data/definitions/328.html", "https://cwe.mitre.org/data/definitions/359.html"]
  },
  "a4": {
    'Name': 'XML External Entities (XXE)', 
    'Description': 'The software processes an XML document that can contain XML entities with URIs that resolve to documents outside of the intended sphere of control, causing the product to embed incorrect documents into its output.', 
    'Image': 'a4.png',
    'Possible mitigations': {
      "Phase: Implementation" : [
        "•  Before parsing XML files with associated DTDs (Document Type Definition), scan for recursive entity declarations and do not continue parsing potentially explosive content.",
        "•  Many XML parsers and validators can be configured to disable external entity expansion."
      ],
      "Phase: Operation" : [
        "•  If possible, prohibit the use of DTDs or use an XML parser that limits the expansion of recursive DTD entities."
        ]
    },
    'Reference links': ["https://cwe.mitre.org/data/definitions/611.html", "https://cwe.mitre.org/data/definitions/776.html"]
  },
  "a5": {
    'Name': 'Broken Access Control', 
    'Description': 'The software does not restrict or incorrectly restricts access to a resource from an unauthorized actor.', 
    'Image': 'a5.png',
      'Possible mitigations': {
        "Phase: Arquitecture and design" : [
          "•  Recommendations include restricting access to important directories or files by adopting a need to know requirement for both the document and server root, and turning off features such as Automatic Directory Listings that could expose private files and provide information that could be utilized by an attacker when formulating or conducting an attack."
          ],
        "Phase: Implementation" : [
          "•  Handle exceptions internally and do not display errors containing potentially sensitive information to a user.",
          "•  Ensure that error messages only contain minimal details that are useful to the intended audience, and nobody else. The messages need to strike the balance between being too cryptic and not being cryptic enough. They should not necessarily reveal the methods that were used to determine the error. Such detailed information can be used to refine the original attack to increase the chances of success. If errors must be tracked in some detail, capture them in log messages - but consider what could occur if the log messages can be viewed by attackers. Avoid recording highly sensitive information such as passwords in any form. Avoid inconsistent messaging that might accidentally tip off an attacker about internal state, such as whether a username is valid or not."
        ],
        "Phase: System Configuration" : [
          "•  Create default error pages or messages that do not leak any information."
          ]
      },
      'Reference links': ["https://cwe.mitre.org/data/definitions/209.html", "https://cwe.mitre.org/data/definitions/508.html"]
    },
  "a6": {
    'Name': 'Security Misconfiguration', 
    'Description': 'The software generates an error message that includes sensitive information about its environment, users, or associated data.', 
    'Image': 'a6.png',
    'Possible mitigations': {
      "Phase: Arquitecture and design" : [
        "•  Recommendations include restricting access to important directories or files by adopting a need to know requirement for both the document and server root, and turning off features such as Automatic Directory Listings that could expose private files and provide information that could be utilized by an attacker when formulating or conducting an attack."
        ],
      "Phase: Implementation" : [
        "•  Handle exceptions internally and do not display errors containing potentially sensitive information to a user.",
        "•  Ensure that error messages only contain minimal details that are useful to the intended audience, and nobody else. The messages need to strike the balance between being too cryptic and not being cryptic enough. They should not necessarily reveal the methods that were used to determine the error. Such detailed information can be used to refine the original attack to increase the chances of success. If errors must be tracked in some detail, capture them in log messages - but consider what could occur if the log messages can be viewed by attackers. Avoid recording highly sensitive information such as passwords in any form. Avoid inconsistent messaging that might accidentally tip off an attacker about internal state, such as whether a username is valid or not."
      ],
      "Phase: System Configuration" : [
        "•  Create default error pages or messages that do not leak any information."
        ]
    },
    'Reference links': ["https://cwe.mitre.org/data/definitions/209.html", "https://cwe.mitre.org/data/definitions/508.html"]
  },
  "a7": {
    'Name': 'Cross-Site Scripting (XSS)', 
    'Description': 'The software does not neutralize or incorrectly neutralizes user-controllable input before it is placed in output that is used as a web page that is served to other users.', 
    'Image': 'a7.png',
    'Possible mitigations': {
      "Phase: Arquitecture and design" : [
        "•  Libraries or Frameworks: Use a vetted library or framework that does not allow this weakness to occur or provides constructs that make this weakness easier to avoid. Examples of libraries and frameworks that make it easier to generate properly encoded output include Microsoft's Anti-XSS library, the OWASP ESAPI Encoding module, and Apache Wicket.",
        "•  For any security checks that are performed on the client side, ensure that these checks are duplicated on the server side, in order to avoid CWE-602. Attackers can bypass the client-side checks by modifying values after the checks have been performed, or by changing the client to remove the client-side checks entirely. Then, these modified values would be submitted to the server.",
        "•  Attack Surface Reduction: Understand all the potential areas where untrusted inputs can enter your software: parameters or arguments, cookies, anything read from the network, environment variables, reverse DNS lookups, query results, request headers, URL components, e-mail, files, filenames, databases, and any external systems that provide data to the application. Remember that such inputs may be obtained indirectly through API calls. This technique can be helpful when it is possible to store client state and sensitive information on the server side instead of in cookies, headers, hidden form fields, etc."
      ],
      "Phase: Implementation" : [
        '•  Input Validation: Assume all input is malicious. Use an "accept known good" input validation strategy, i.e., use a list of acceptable inputs that strictly conform to specifications. Reject any input that does not strictly conform to specifications, or transform it into something that does.',
        "•  Attack Surface Reduction: To help mitigate XSS attacks against the user's session cookie, set the session cookie to be HttpOnly. In browsers that support the HttpOnly feature (such as more recent versions of Internet Explorer and Firefox), this attribute can prevent the user's session cookie from being accessible to malicious client-side scripts that use document.cookie. This is not a complete solution, since HttpOnly is not supported by all browsers. More importantly, XMLHTTPRequest and other powerful browser technologies provide read access to HTTP headers, including the Set-Cookie header in which the HttpOnly flag is set.",
        "•  Output Encoding: Use and specify an output encoding that can be handled by the downstream component that is reading the output. Common encodings include ISO-8859-1, UTF-7, and UTF-8. When an encoding is not specified, a downstream component may choose a different encoding, either by assuming a default encoding or automatically inferring which encoding is being used, which can be erroneous. When the encodings are inconsistent, the downstream component might treat some character or byte sequences as special, even if they are not special in the original encoding. Attackers might then be able to exploit this discrepancy and conduct injection attacks; they even might be able to bypass protection mechanisms that assume the original encoding is also being used by the downstream component. The problem of inconsistent output encodings often arises in web pages. If an encoding is not specified in an HTTP header, web browsers often guess about which encoding is being used. This can open up the browser to subtle XSS attacks."
      ]
    },
    'Reference links': ["https://cwe.mitre.org/data/definitions/79.html"]
  },
  "a8": {
    'Name': 'Insecure Deserialization', 
    'Description': 'The application deserializes untrusted data without sufficiently verifying that the resulting data will be valid.', 
    'Image': 'a8.png',
      'Possible mitigations': {
        "Phase: Arquitecture and design" : [
          "•  If available, use the signing/sealing features of the programming language to assure that deserialized data has not been tainted. For example, a hash-based message authentication code (HMAC) could be used to ensure that data has not been modified.",
          "•  Make fields transient to protect them from deserialization. An attempt to serialize and then deserialize a class containing transient fields will result in NULLs where the transient data should be. This is an excellent way to prevent time, environment-based, or sensitive variables from being carried over and used improperly."
          ],
        "Phase: Implementation" : [
          "•  When deserializing data, populate a new object rather than just deserializing. The result is that the data flows through safe input validation and that the functions are safe.",
          "•  Explicitly define a final object() to prevent deserialization.",
          "•  Avoid having unnecessary types or gadgets available that can be leveraged for malicious ends. This limits the potential for unintended or unauthorized types and gadgets to be leveraged by the attacker. Whitelist acceptable classes. Note: new gadgets are constantly being discovered, so this alone is not a sufficient mitigation."
        ],
        "Phase: System Configuration" : [
          "•  Create default error pages or messages that do not leak any information."
          ]
      },
      'Reference links': ["https://cwe.mitre.org/data/definitions/502.html"]
    },
  "a9": {
    'Name': 'Using Components with Known Vulnerabilities', 
    'Description': 'Since "known vulnerabilities" can arise from any kind of weakness, it is not possible to map this OWASP category to other CWE entries, since it would effectively require mapping this category to ALL weaknesses.', 
    'Image': 'a9.png',
      'Possible mitigations': {
        "General" : [
          "•  Check for known vulnerabilities in the reference links."
          ],  
      },
      'Reference links': ["https://cwe.mitre.org/data/definitions/1035.html", "https://cwe.mitre.org/data/definitions/1026.html"]
    },
  "a10": {
    'Name': 'Insufficient Logging & Monitoring', 
    'Description': 'The application does not record or display information that would be important for identifying the source or nature of an attack, or determining if an action is safe.', 
    'Image': 'a10.png',
      'Possible mitigations': {
        "Phase: Arquitecture and design" : [
          "•  Use a centralized logging mechanism that supports multiple levels of detail. Ensure that all security-related successes and failures can be logged."
          ],
        "Phase: Operation" : [
          "•  Be sure to set the level of logging appropriately in a production environment. Sufficient data should be logged to enable system administrators to detect attacks, diagnose errors, and recover from attacks. At the same time, logging too much data (CWE-779) can cause the same problems."
          ]
      },
      'Reference links': ["https://cwe.mitre.org/data/definitions/223.html", "https://cwe.mitre.org/data/definitions/778.html"]
    },

}

#Aux function that cleans up the code
def eliminate_labels(string):
  string_copy = []
  final_string = ""
  x1 = string.find("<")
  x2 = string.find(">")

  for i in range(0,x1):
    string_copy.append(string[i])
  
  for i in range(x2+1,len(string)):
    string_copy.append(string[i])
  
  for i in range(0,len(string_copy)):
    final_string = final_string + str(string_copy[i])

  return final_string

#Aux function that cleans up the code
def clean_code(string):
  string.strip()
  while (string.find("<") != -1):
    string = eliminate_labels(string)
  return string



#######################################################################
###################    Document report     ############################
#######################################################################

os.system("pip3 install python-docx")
os.system("pip install matplotlib")

#-------------------------------------------------------------------------------------
import docx
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import matplotlib.pyplot as plt 
from datetime import date

today = date.today()
day = today.strftime("%d/%m/%Y")

document = Document()

########################## Styles #########################################

styles = document.styles
style = styles.add_style('mytitle', WD_STYLE_TYPE.PARAGRAPH)
style.font.size = Pt(40)
style.font.color.rgb = RGBColor(47, 84, 150)

style = styles.add_style('myversion', WD_STYLE_TYPE.PARAGRAPH)
style.font.size = Pt(16)
style.font.color.rgb = RGBColor(0, 176, 240)

style = styles.add_style('mysubtitle', WD_STYLE_TYPE.PARAGRAPH)
style.font.size = Pt(24)
style.font.color.rgb = RGBColor(68, 114, 196)

style = styles.add_style('mydate', WD_STYLE_TYPE.PARAGRAPH)
style.font.size = Pt(16)


######################### Title page ########################################

document.add_paragraph()
document.add_paragraph()
p = document.add_paragraph(style = 'mytitle')
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(project_name.upper())
run.bold = True
run.italic = True

document.add_paragraph()
p = document.add_paragraph(style = 'myversion')
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("VERSION "+ projectVersion.upper())
run.italic = True

document.add_paragraph()
document.add_paragraph()
document.add_paragraph()
p = document.add_paragraph(style = 'mysubtitle')
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Code Analysis")
run.bold = True

document.add_paragraph()
document.add_paragraph()
document.add_paragraph()
document.add_paragraph()
document.add_paragraph()
document.add_paragraph()
p = document.add_paragraph(style = 'mydate')
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run(str(day))
run.bold = True


#############################################################################


p = document.add_heading('Detected vulnerabilities OwaspTop10', level=0)
p.paragraph_format.page_break_before = True



############ Creates types of values (a1, a2...) ############
if facets != []:
    #values = ["a1", "a4",...]
    #counts = [ 6, 3,...]
    values = []
    counts = []
    for i in range(0,len(facets[0]['values'])):
        p = document.add_paragraph()
        value = facets[0]['values'][i]['val']
        count = facets[0]['values'][i]['count']
        values.append(value)
        counts.append(count)
        run = "{0} ({1}): ".format(vulnerabilities[value]["Name"],values[i].upper())
        p.add_run(run).bold = True
        run = "{0}".format(facets[0]['values'][i]['count'])
        p.add_run(run)

    ###################### Plot #################################

    document.add_paragraph()
      
    # color for each label 
    color_example = ['r', 'g', 'y', 'b', 'r', 'g', 'y', 'b', 'r', 'g'] 
    colors = []
    explode = []
    for i in range(0,len(values)):
        colors.append(color_example[i])
        explode.append(0.1)

    # plotting the pie chart 
    plt.pie(counts, labels = values, colors = colors,  
            startangle=90, shadow = True, explode = explode, 
            radius = 1.2, autopct = '%1.1f%%') 

    plt.savefig('plot.png', bbox_inches = 'tight')

    document.add_picture('plot.png')
    os.system("DEL /F /A plot.png")

    #############################################################################


############ For each value, display description, possible mitigations, image... ############
    for v in range(0,len(values)):
        document.add_paragraph()
        #####AX - Code Injection #######
        p = document.add_heading(values[v].upper() + " - " + vulnerabilities[values[v]]["Name"], level=0)
        p.paragraph_format.page_break_before = True
        document.add_heading('Description', level=1)
        document.add_paragraph()
        p = document.add_paragraph()
        p.add_run(vulnerabilities[values[v]]['Description'])
        document.add_paragraph()

        ####### Show vulnerabilities in the code #######
        document.add_heading('Detected vulnerabilities in code', level=1)
        for i in range(0,len(filtered_issues[values[v]])):
            document.add_paragraph("Vulnerability {0} detected at".format(i+1), style='Intense Quote')
            
            p = document.add_paragraph()
            p.add_run("{0}: {1}".format("Component ", filtered_issues[values[v]][i]['component']))

            p = document.add_paragraph()
            p.add_run("{0}: {1}".format("Starts at line ", filtered_issues[values[v]][i]['textRange']['startLine']))

            p = document.add_paragraph()
            p.add_run("{0}: {1}".format("Ends at line ", filtered_issues[values[v]][i]['textRange']['endLine']))
            
            #"Python-Proyect-recursive_IDS/du2.py"
            #"test1:/nbody/Analizar/distributed/du_files/du2.py"
            #http://localhost:9000/api/sources/show?key=Python-Proyect-recursive_IDS:Analizar/du_1.py
            
            #URL = "http://localhost:9000/api/sources/show?key=Python-Proyect-recursive_IDS:Analizar/du_1.py"
            #URL = "http://localhost:9000/api/sources/show?key=" + filtered_issues[values[v]][i]['component']

            URL = "http://localhost:9000/api/sources/show" 
            PARAMS = {'key': filtered_issues[values[v]][i]['component']}
            r = requests.get(url = URL, params = PARAMS)

            data = r.json()

            lines = data['sources']
            for l in range(filtered_issues[values[v]][i]['textRange']['startLine'] - 1, filtered_issues[values[v]][i]['textRange']['endLine']):
                print(lines[l])
                p = document.add_paragraph()
                code = lines[l][1]
                code = clean_code(code)
                p.add_run("{0}: {1}".format(lines[l][0], code)).bold = True



        ############## Mitigations ##############
        p = document.add_heading('Possible mitigations', level=1)
        p.paragraph_format.page_break_before = True
        for phase in vulnerabilities[values[v]]['Possible mitigations']:
          document.add_paragraph()
          p = document.add_paragraph()
          p.add_run(phase).bold = True
          for mitigation in vulnerabilities[values[v]]['Possible mitigations'][phase]:
            p = document.add_paragraph()
            p.add_run(mitigation)

        ############## Links ################
        document.add_heading('Reference links', level=1)
        document.add_paragraph()
        for link in vulnerabilities[values[v]]['Reference links']:
          p = document.add_paragraph()
          p.add_run(link).bold = True

document.save('Analysis report.docx')


